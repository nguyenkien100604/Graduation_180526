# -*- coding: utf-8 -*-
"""Đọc file đa định dạng (CSV, JSON, XLSX, DOCX, TXT) → DataFrame."""

from __future__ import annotations

import csv
import io
import json
from pathlib import Path
from typing import Any, BinaryIO

import pandas as pd

SUPPORTED_EXTENSIONS = {".csv", ".json", ".xlsx", ".xls", ".docx", ".txt", ".tsv"}


def detect_format(filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Định dạng '{ext}' chưa hỗ trợ. Hỗ trợ: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )
    return ext


def _read_bytes(source: str | Path | BinaryIO | bytes) -> bytes:
    if isinstance(source, bytes):
        return source
    if isinstance(source, (str, Path)):
        return Path(source).read_bytes()
    return source.read()


def _read_csv_bytes(data: bytes) -> pd.DataFrame:
    for enc in ("utf-8", "utf-8-sig", "cp1252", "latin1"):
        try:
            return pd.read_csv(io.BytesIO(data), encoding=enc)
        except UnicodeDecodeError:
            continue
    return pd.read_csv(io.BytesIO(data), encoding="utf-8", errors="replace")


def _read_json_bytes(data: bytes) -> pd.DataFrame:
    text = data.decode("utf-8", errors="replace").strip()
    obj = json.loads(text)
    if isinstance(obj, list):
        return pd.json_normalize(obj)
    if isinstance(obj, dict):
        for key in ("data", "records", "transactions", "items", "results"):
            if key in obj and isinstance(obj[key], list):
                return pd.json_normalize(obj[key])
        return pd.json_normalize(obj)
    raise ValueError("JSON phải là mảng bản ghi hoặc object chứa mảng (data/records/...).")


def _read_xlsx_bytes(data: bytes, sheet_name: str | int | None = None) -> pd.DataFrame:
    bio = io.BytesIO(data)
    name = sheet_name if sheet_name is not None else 0
    return pd.read_excel(bio, sheet_name=name, engine="openpyxl")


def _read_txt_bytes(data: bytes) -> pd.DataFrame:
    text = data.decode("utf-8", errors="replace").strip()
    if not text:
        raise ValueError("File TXT trống.")

    if text.startswith("[") or text.startswith("{"):
        return _read_json_bytes(data)

    # CSV/TSV sniff
    try:
        dialect = csv.Sniffer().sniff(text[:4096], delimiters=",;\t|")
        delim = dialect.delimiter
    except csv.Error:
        delim = "\t" if text.count("\t") > text.count(",") else ","

    return pd.read_csv(io.StringIO(text), sep=delim)


def _read_docx_bytes(data: bytes) -> pd.DataFrame:
    try:
        from docx import Document
    except ImportError as e:
        raise ImportError("Cài python-docx: pip install python-docx") from e

    doc = Document(io.BytesIO(data))
    best: pd.DataFrame | None = None

    for table in doc.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if len(rows) < 2:
            continue
        header = rows[0]
        body = rows[1:]
        if not any(header):
            continue
        df = pd.DataFrame(body, columns=header)
        if best is None or len(df) > len(best):
            best = df

    if best is not None and not best.empty:
        return best

    # Không có bảng: gom paragraph có dấu phân cách
    lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if not lines:
        raise ValueError("DOCX không có bảng hoặc đoạn văn có dữ liệu.")

    block = "\n".join(lines)
    try:
        dialect = csv.Sniffer().sniff(block[:4096], delimiters=",;\t|")
        delim = dialect.delimiter
    except csv.Error:
        delim = "\t" if block.count("\t") > block.count(",") else ","
    return pd.read_csv(io.StringIO(block), sep=delim)


def _parse_free_text(text: str) -> pd.DataFrame:
    """Parser cục bộ — không tốn quota API."""
    from text_parser import extract_transactions_from_text

    return extract_transactions_from_text(text)


def ai_extract_table_from_text(text: str) -> pd.DataFrame | None:
    """Groq/Gemini; trả None nếu hết quota (caller dùng parser cục bộ)."""
    if len(text) < 20:
        return None

    from ai_client import AIQuotaError, ai_available, ai_json_array
    from ingest_core import COLUMNS

    if not ai_available():
        return None

    try:
        rows = ai_json_array(
            "Extract banking transaction rows from text. "
            "Return a JSON array of objects only. Use English keys close to: "
            + ", ".join(COLUMNS),
            text[:12000],
        )
        if rows:
            return pd.json_normalize(rows)
    except AIQuotaError:
        return None
    return None


def _docx_plain_text(data: bytes) -> str:
    from docx import Document

    return "\n".join(p.text for p in Document(io.BytesIO(data)).paragraphs if p.text.strip())


def _load_docx_bytes(data: bytes, *, use_ai_extract: bool) -> tuple[pd.DataFrame, dict[str, Any]]:
    """DOCX: bảng → CSV trong đoạn → parser văn bản → AI (tùy chọn)."""
    meta: dict[str, Any] = {}
    try:
        return _read_docx_bytes(data), meta
    except (ValueError, pd.errors.ParserError):
        pass

    text = _docx_plain_text(data)
    if not text:
        raise ValueError("DOCX không có nội dung.")

    df = _parse_free_text(text)
    if not df.empty:
        meta["extract"] = "text_parser"
        return df, meta

    if use_ai_extract:
        df = ai_extract_table_from_text(text)
        if df is not None and not df.empty:
            meta["extract"] = "ai"
            return df, meta

    raise ValueError(
        "DOCX dạng văn bản: không parse được. Bật «AI trích xuất» (cần quota Gemini) "
        "hoặc dùng file có bảng / CSV / JSON."
    )


def load_file(
    source: str | Path | BinaryIO | bytes,
    filename: str,
    *,
    sheet_name: str | int | None = None,
    use_ai_extract: bool = False,
) -> tuple[pd.DataFrame, dict[str, Any]]:
    """
    Đọc file bất kỳ định dạng hỗ trợ.
    Returns (DataFrame thô, metadata).
    """
    ext = detect_format(filename)
    data = _read_bytes(source)

    meta: dict[str, Any] = {"format": ext, "filename": filename}

    if ext == ".csv" or ext == ".tsv":
        df = _read_csv_bytes(data)
    elif ext == ".json":
        df = _read_json_bytes(data)
    elif ext in (".xlsx", ".xls"):
        df = _read_xlsx_bytes(data, sheet_name=sheet_name)
        meta["sheet"] = sheet_name if sheet_name is not None else 0
    elif ext == ".txt":
        df = _read_txt_bytes(data)
    elif ext == ".docx":
        df, extra = _load_docx_bytes(data, use_ai_extract=use_ai_extract)
        meta.update(extra)
    else:
        raise ValueError(f"Unsupported: {ext}")

    if df is None or df.empty:
        if ext == ".txt":
            text = data.decode("utf-8", errors="replace")
            df = _parse_free_text(text)
            if df.empty and use_ai_extract:
                df = ai_extract_table_from_text(text) or df
                meta["extract"] = "ai" if not df.empty else meta.get("extract")
            if df.empty:
                raise ValueError("TXT: không đọc được dữ liệu có cấu trúc.")
        else:
            raise ValueError("File không có dòng dữ liệu.")

    meta["rows"] = len(df)
    meta["columns"] = list(df.columns)
    return df, meta
